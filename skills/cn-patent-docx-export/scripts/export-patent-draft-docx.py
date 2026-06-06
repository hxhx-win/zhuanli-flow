#!/usr/bin/env python3
"""Export patent Markdown to DOCX using pandoc (Linux) or Word COM (Windows)."""
import argparse
import copy
import json
import shutil
import subprocess
import sys
import platform
from pathlib import Path


import re
import tempfile

EXPORT_SECTION_KEYS = ['摘要', '摘要附图', '权利要求书', '说明书']
OPTIONAL_EXPORT_SECTION_KEYS = ['说明书附图']
SPEC_SUBSECTION_TITLES = {'发明名称', '技术领域', '背景技术', '发明内容', '附图说明', '具体实施方式'}
INTERNAL_SUBHEADING_RE = re.compile(r'^(?:实施例[0-9一二三四五六七八九十]+|可选实施方式|替代实施方式)$')
CNIPA_PAGE_SIZE = {'w': '11906', 'h': '16838'}
CNIPA_PAGE_MARGINS = {
    'top': '1418',
    'left': '1418',
    'right': '850',
    'bottom': '850',
    'header': '567',
    'footer': '567',
}
CNIPA_BODY_FONT_EAST_ASIA = '宋体'
CNIPA_BODY_FONT_ASCII = 'Times New Roman'
CNIPA_BODY_FONT_SIZE = '28'
CNIPA_HEADER_FONT_SIZE = '28'
CNIPA_HEADER_CHAR_SPACING = '280'  # 14pt 字号下一个字符宽度 (twentieths of a point)
CNIPA_LINE_SPACING = '360'
CNIPA_FIRST_LINE_CHARS = '200'
CNIPA_FIRST_LINE_TWIPS = '560'
CNIPA_ABSTRACT_MAX_CHARS = 300
CLAIM_FORBIDDEN_EXPRESSIONS = (
    '如说明书所述',
    '如上所述',
    '参见说明书',
    '见说明书',
    '说明书中所述',
)


def normalize_figure_number_label(entry):
    """Return the official below-figure label: figure number only, no title."""
    raw = str(entry.get('figureNumber', '')).strip()
    if not raw:
        raw = str(entry.get('caption', '')).strip()
    match = re.match(r'^图\s*([0-9]+)', raw)
    if match:
        return f"图{match.group(1)}"
    match = re.match(r'^([0-9]+)$', raw)
    if match:
        return f"图{match.group(1)}"
    return ''


def get_markdown_section(text, title):
    """Return the body of a top-level patent Markdown section."""
    pattern = re.compile(rf'^##\s+{re.escape(title)}\s*$\n?(.*?)(?=^##\s+|\Z)', re.MULTILINE | re.DOTALL)
    match = pattern.search(text)
    return match.group(1).strip() if match else ''


def split_claim_paragraphs(claim_body):
    """Extract numbered claim paragraphs from the claims section."""
    claims = []
    current = None
    for line in claim_body.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        match = re.match(r'^([0-9]+)[\.、]\s*(.*)', stripped)
        if match:
            if current:
                claims.append(current)
            current = match.group(2).strip()
        elif current:
            current += stripped
    if current:
        claims.append(current)
    return claims


def preflight_check():
    """Validate platform-specific dependencies before export."""
    errors = []
    warnings = []
    if platform.system() == "Windows":
        if not shutil.which('pandoc'):
            try:
                import win32com.client  # noqa: F401
            except ImportError:
                errors.append("Windows 上需 pandoc 或 win32com 其一 "
                              "(winget install --id JohnMacFarlane.Pandoc / pip install pywin32)")
    else:
        if not shutil.which('pandoc'):
            errors.append("pandoc not found (apt install pandoc / brew install pandoc)")
    try:
        import docx  # noqa: F401
    except ImportError:
        errors.append("python-docx not available (pip install python-docx)")
    if warnings:
        for w in warnings:
            print(f"WARNING: {w}", file=sys.stderr)
    if errors:
        print("BLOCKED: Export preflight failed:", file=sys.stderr)
        for e in errors:
            print(f"  - {e}", file=sys.stderr)
        print("Run patent-env-check script for full capability matrix.", file=sys.stderr)
        sys.exit(1)



def validate_markdown_source(input_path):
    """Pre-pandoc Markdown source validation. Checks section keys, order, formula format,
    and residual internal-only content that would break the export."""
    text = Path(input_path).read_text(encoding='utf-8')
    lines = text.split('\n')
    errors = []
    warnings = []

    # 1. Section key matching — script expects exact ## titles
    expected_keys = set(EXPORT_SECTION_KEYS)
    h2_titles = []
    for line in lines:
        if line.startswith('## ') and not line.startswith('### '):
            title = line[3:].strip()
            h2_titles.append(title)

    found_keys = set(h2_titles) & expected_keys
    missing_keys = expected_keys - found_keys
    if missing_keys:
        errors.append(f"缺少必要 ## 标题: {missing_keys}。脚本按精确 key 匹配插入 section break，"
                      f"'说明书摘要' 应为 '摘要'，需要独立 '## 摘要附图' 标题。")

    # 2. Section order
    expected_order = EXPORT_SECTION_KEYS
    found_order = [t for t in h2_titles if t in expected_keys]
    if found_order != expected_order:
        errors.append(f"Section 顺序错误: 当前 {found_order}，期望 {expected_order}")

    invalid_h2 = [t for t in h2_titles if t in SPEC_SUBSECTION_TITLES or INTERNAL_SUBHEADING_RE.match(t)]
    if invalid_h2:
        errors.append(f"说明书内部小标题不能使用 ##: {invalid_h2}。请放在 ## 说明书 下并改用 ###。")

    # 3. Pseudo-formula text
    pseudo_re = re.compile(r'(?<!\$)(?:sum|sqrt|mu|sigma|alpha|beta)\s*\(', re.IGNORECASE)
    for i, line in enumerate(lines, 1):
        if pseudo_re.search(line) and '$' not in line and '$$' not in line:
            errors.append(f"第{i}行疑似伪公式文本（应使用 LaTeX $..$ 或 $$..$$）: {line.strip()[:60]}")

    # 4. Standalone formula uses $$
    dollar_lines = [i for i, l in enumerate(lines, 1) if l.strip().startswith('$$')]
    if len(dollar_lines) % 2 != 0:
        errors.append(f"独立公式 $$ 边界行数为奇数({len(dollar_lines)})，存在未闭合公式")

    # 5. Residual internal sections
    for i, line in enumerate(lines, 1):
        if re.match(r'^#+\s*(证据来源|风险点|待确认事项)', line):
            errors.append(f"第{i}行残留内部章节（不应进入正式稿）: {line.strip()}")

    # 6. 附图交接清单 residual (would cause Step 5b to delete appended figures)
    for i, line in enumerate(lines, 1):
        if '附图交接清单' in line:
            errors.append(f"第{i}行残留附图交接清单（会导致导出脚本误删已嵌入附图）")

    # 7. Specification subheading paragraphs must stand alone.
    spec_subsections = SPEC_SUBSECTION_TITLES
    for i, line in enumerate(lines, 1):
        m = re.match(r'^\*\*(.+?)\*\*\s*$', line.strip())
        if m and m.group(1) in spec_subsections:
            if i < len(lines) and lines[i].strip():
                errors.append(f"第{i}行说明书小标题后必须空一行，否则会与正文合并并产生首行缩进: {line.strip()}")

    # 8. CNIPA abstract constraints.
    abstract_body = get_markdown_section(text, '摘要')
    if abstract_body:
        abstract_plain = re.sub(r'[\s#*_`|]', '', abstract_body)
        if len(abstract_plain) > CNIPA_ABSTRACT_MAX_CHARS:
            errors.append(f"摘要文字超过{CNIPA_ABSTRACT_MAX_CHARS}字: 当前约{len(abstract_plain)}字")
        if re.search(r'^#{1,6}\s+', abstract_body, re.MULTILINE):
            errors.append("摘要正文不得包含 Markdown 小标题")
        if '|' in abstract_body or re.search(r'!\[|<image\b', abstract_body):
            errors.append("摘要正文不得包含表格或图片")

    # 9. CNIPA claim constraints.
    claim_body = get_markdown_section(text, '权利要求书')
    if claim_body:
        if re.search(r'!\[|<image\b|<whiteboard\b', claim_body):
            errors.append("权利要求书不得包含插图、图片或画板")
        for expression in CLAIM_FORBIDDEN_EXPRESSIONS:
            if expression in claim_body:
                errors.append(f"权利要求书不得使用引用式表述: {expression}")
        for idx, claim in enumerate(split_claim_paragraphs(claim_body), 1):
            normalized = claim.strip()
            if not normalized.endswith('。'):
                errors.append(f"权利要求{idx}必须以句号结尾")
            internal = normalized[:-1]
            if '。' in internal:
                errors.append(f"权利要求{idx}仅允许在结尾处使用句号")

    # 10. Specification required parts.
    spec_body = get_markdown_section(text, '说明书')
    if spec_body:
        required_parts = ['技术领域', '背景技术', '发明内容', '具体实施方式']
        for part in required_parts:
            if not re.search(rf'^###\s+{re.escape(part)}\s*$', spec_body, re.MULTILINE):
                errors.append(f"说明书缺少官方要求部分: {part}")

    # Report
    if warnings:
        for w in warnings:
            print(f"[md-validate] WARN: {w}", file=sys.stderr)
    if errors:
        print("[md-validate] BLOCKED: Markdown 源文件预检失败:", file=sys.stderr)
        for e in errors:
            print(f"  - {e}", file=sys.stderr)
        print("[md-validate] 请修正后重新导出。", file=sys.stderr)
        sys.exit(1)
    else:
        print("[md-validate] Markdown 源文件预检通过", file=sys.stderr)


def check_pandoc():
    if not shutil.which('pandoc'):
        print("ERROR: pandoc not found.", file=sys.stderr)
        print("BLOCKED: DOCX export requires pandoc. Do NOT fall back to plain text.", file=sys.stderr)
        sys.exit(1)


NON_PATENT_SECTIONS = re.compile(
    r'^##\s+(?:附图[（(]|证据来源|风险点|待确认|参考文献|备注|工作记录|质量检查)',
    re.MULTILINE
)


def truncate_non_patent_content(text):
    """Remove non-patent sections (evidence, risks, mermaid drafts) from the end."""
    m = NON_PATENT_SECTIONS.search(text)
    if m:
        return text[:m.start()].rstrip()
    return text


def preprocess_latex_fences(text):
    """Convert ```latex ... ``` fences to $$...$$ display math for pandoc."""
    lines = text.split('\n')
    result = []
    in_latex = False
    latex_buf = []
    for line in lines:
        stripped = line.strip()
        if not in_latex and re.match(r'^```latex\s*$', stripped):
            in_latex = True
            latex_buf = []
            continue
        if in_latex:
            if stripped == '```':
                in_latex = False
                result.append('')
                result.append('$$')
                result.append('\n'.join(latex_buf))
                result.append('$$')
                result.append('')
            else:
                latex_buf.append(line)
            continue
        result.append(line)
    return '\n'.join(result)


def export_with_pandoc(input_path, output_path, reference_doc=None):
    text = Path(input_path).read_text(encoding='utf-8')
    text = truncate_non_patent_content(text)
    processed = preprocess_latex_fences(text)
    processed = re.sub(r'\\\((.+?)\\\)', r'$\1$', processed)

    tmp = tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8')
    tmp.write(processed)
    tmp.close()

    cmd = ['pandoc', tmp.name, '-o', str(output_path),
           '--from', 'markdown+tex_math_dollars+raw_tex',
           '--wrap', 'none']
    if reference_doc and Path(reference_doc).exists():
        cmd.extend(['--reference-doc', str(reference_doc)])
    result = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8", errors="replace")
    Path(tmp.name).unlink(missing_ok=True)
    if result.returncode != 0:
        print(f"ERROR: pandoc failed: {result.stderr}", file=sys.stderr)
        sys.exit(1)


SECTION_HEADERS = {
    '摘要': '说明书摘要',
    '摘要附图': '摘要附图',
    '权利要求书': '权利要求书',
    '说明书': '说明书',
    '说明书附图': '说明书附图',
}

SECTION_ORDER = EXPORT_SECTION_KEYS + OPTIONAL_EXPORT_SECTION_KEYS


def make_header_part(doc, header_text):
    """Create a header part with centered 16pt 宋体 text and bottom border."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    header_part, rId = doc.part.add_header_part()
    hdr_elem = header_part.element
    for child in list(hdr_elem):
        hdr_elem.remove(child)
    hp = OxmlElement('w:p')
    hpPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    hpPr.append(jc)
    outline = OxmlElement('w:outlineLvl')
    outline.set(qn('w:val'), '0')
    hpPr.append(outline)
    hp_rPr = OxmlElement('w:rPr')
    hp_fonts = OxmlElement('w:rFonts')
    hp_fonts.set(qn('w:ascii'), '黑体')
    hp_fonts.set(qn('w:eastAsia'), '黑体')
    hp_rPr.append(hp_fonts)
    hp_bold = OxmlElement('w:b')
    hp_rPr.append(hp_bold)
    hp_spacing = OxmlElement('w:spacing')
    hp_spacing.set(qn('w:val'), CNIPA_HEADER_CHAR_SPACING)
    hp_rPr.append(hp_spacing)
    hp_sz = OxmlElement('w:sz')
    hp_sz.set(qn('w:val'), CNIPA_HEADER_FONT_SIZE)
    hp_rPr.append(hp_sz)
    hp.append(hpPr)
    hpPr.append(hp_rPr)
    hr = OxmlElement('w:r')
    hrPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), CNIPA_HEADER_FONT_SIZE)
    hrPr.append(sz)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), '黑体')
    rFonts.set(qn('w:eastAsia'), '黑体')
    hrPr.append(rFonts)
    b = OxmlElement('w:b')
    hrPr.append(b)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), CNIPA_HEADER_CHAR_SPACING)
    hrPr.append(spacing)
    hr.append(hrPr)
    ht = OxmlElement('w:t')
    ht.text = header_text
    ht.set(qn('xml:space'), 'preserve')
    hr.append(ht)
    hp.append(hr)
    hdr_elem.append(hp)
    return rId


def make_footer_part(doc):
    """Create a footer part with centered PAGE field."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    footer_part, rId = doc.part.add_footer_part()
    ftr_elem = footer_part.element
    for child in list(ftr_elem):
        ftr_elem.remove(child)

    fp = OxmlElement('w:p')
    fpPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    fpPr.append(jc)
    fp.append(fpPr)

    def add_field_run(field_type, text=None):
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), CNIPA_BODY_FONT_ASCII)
        rFonts.set(qn('w:eastAsia'), CNIPA_BODY_FONT_EAST_ASIA)
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), CNIPA_BODY_FONT_SIZE)
        rPr.append(sz)
        run.append(rPr)
        if field_type == 'begin':
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), 'begin')
            run.append(fld)
        elif field_type == 'instr':
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = ' PAGE '
            run.append(instr)
        elif field_type == 'separate':
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), 'separate')
            run.append(fld)
        elif field_type == 'text':
            t = OxmlElement('w:t')
            t.text = text or '1'
            run.append(t)
        elif field_type == 'end':
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), 'end')
            run.append(fld)
        fp.append(run)

    for field_type in ('begin', 'instr', 'separate', 'text', 'end'):
        add_field_run(field_type)
    ftr_elem.append(fp)
    return rId


def apply_cnipa_section_properties(doc, sectPr, header_text=None):
    """Apply official page settings, optional header/footer, and restarted page numbers."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is None:
        pgSz = OxmlElement('w:pgSz')
        sectPr.append(pgSz)
    pgSz.set(qn('w:w'), CNIPA_PAGE_SIZE['w'])
    pgSz.set(qn('w:h'), CNIPA_PAGE_SIZE['h'])

    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is None:
        pgMar = OxmlElement('w:pgMar')
        sectPr.append(pgMar)
    for key, value in CNIPA_PAGE_MARGINS.items():
        pgMar.set(qn(f'w:{key}'), value)

    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:start'), '1')

    for existing in list(sectPr.findall(qn('w:footerReference'))):
        sectPr.remove(existing)

    if header_text:
        for existing in list(sectPr.findall(qn('w:headerReference'))):
            sectPr.remove(existing)
        header_rid = make_header_part(doc, header_text)
        headerRef = OxmlElement('w:headerReference')
        headerRef.set(qn('w:type'), 'default')
        headerRef.set(qn('r:id'), header_rid)
        sectPr.append(headerRef)

    footer_rid = make_footer_part(doc)
    footerRef = OxmlElement('w:footerReference')
    footerRef.set(qn('w:type'), 'default')
    footerRef.set(qn('r:id'), footer_rid)
    sectPr.append(footerRef)


def make_sectPr_with_header(doc, header_text):
    """Build a sectPr element with page settings and header reference."""
    from docx.oxml import OxmlElement

    sectPr = OxmlElement('w:sectPr')
    apply_cnipa_section_properties(doc, sectPr, header_text)
    return sectPr


def post_process_docx(docx_path, figure_manifest=None):
    """Post-process DOCX: add section headers, remove title paragraphs, convert claims,
    apply paragraph formatting, embed figures, and set page layout."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("[export-patent-draft-docx] python-docx not available, skipping post-processing", file=sys.stderr)
        return

    doc = Document(str(docx_path))
    if len(doc.paragraphs) == 0:
        print("WARNING: Generated DOCX has no paragraphs", file=sys.stderr)
        return

    # --- Step 1: Build sections with correct headers ---
    # Strategy: find section title paragraphs in order. For each title found,
    # the content AFTER it (until the next title) gets that title's header.
    # The H1 title paragraph (first paragraph) is also removed.
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    body = doc.element.body
    all_p_elems = list(body.iterchildren(qn('w:p')))

    # Find title paragraph positions
    title_positions = []  # (index_in_all_p, section_key)
    h1_to_remove = None
    for i, p_elem in enumerate(all_p_elems):
        texts = [t.text for t in p_elem.iter(qn('w:t')) if t.text]
        text = ''.join(texts).strip()
        if text in SECTION_HEADERS:
            title_positions.append((i, text))
        elif i == 0:
            # First paragraph is likely the H1 title - check if it has Heading1 style
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None and 'Heading' in pStyle.get(qn('w:val'), ''):
                    h1_to_remove = p_elem

    # Insert section breaks: for sections 1..N-1, put sectPr in the last paragraph
    # before each title. The sectPr carries the header for the PRECEDING section.
    # Section N's header goes into the body sectPr.
    sections_added = 0
    paras_to_remove = []

    for idx, (pos, key) in enumerate(title_positions):
        header_text = SECTION_HEADERS[key]
        paras_to_remove.append(all_p_elems[pos])

        if idx == 0:
            # First section: find the last paragraph before this title and attach
            # sectPr with the FIRST section's header (说明书摘要)
            # But wait - the first section's header should be on the body's first content.
            # We need to put sectPr on the paragraph BEFORE the NEXT title to end section 1.
            # Actually: sectPr in a paragraph's pPr means "this paragraph is the LAST
            # paragraph of a section". So to give section 1 its header, we put sectPr
            # in the last paragraph of section 1 (= paragraph before title of section 2).
            pass  # handled below

    # Now build the actual section breaks between consecutive titles
    # Section i content: from title_positions[i] to title_positions[i+1]
    # To end section i, put sectPr (with section i's header) in the last para before title[i+1]
    for idx in range(len(title_positions) - 1):
        current_key = title_positions[idx][1]
        next_pos = title_positions[idx + 1][0]
        header_text = SECTION_HEADERS[current_key]

        # Find the last paragraph before next_pos (skipping the title itself)
        last_para_of_section = None
        for j in range(next_pos - 1, -1, -1):
            p = all_p_elems[j]
            if p not in paras_to_remove and p != h1_to_remove:
                last_para_of_section = p
                break

        if last_para_of_section is not None:
            sectPr = make_sectPr_with_header(doc, header_text)
            pPr = last_para_of_section.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                last_para_of_section.insert(0, pPr)
            pPr.append(sectPr)
            sections_added += 1

    # Last section's header goes into body sectPr
    if title_positions:
        last_key = title_positions[-1][1]
        last_header = SECTION_HEADERS[last_key]
        body_sectPr = body.find(qn('w:sectPr'))
        if body_sectPr is None:
            body_sectPr = OxmlElement('w:sectPr')
            body.append(body_sectPr)
        apply_cnipa_section_properties(doc, body_sectPr, last_header)
        sections_added += 1

    # Remove title paragraphs and H1
    if h1_to_remove is not None:
        h1_to_remove.getparent().remove(h1_to_remove)
    for p in paras_to_remove:
        p.getparent().remove(p)

    if sections_added > 0:
        print(f"[post-process] Added {sections_added} section headers", file=sys.stderr)

    # --- Step 3: Convert claim numbering ---
    # Find paragraphs with Heading2 style matching "权利要求\d+" and convert to
    # sequential "N. " prefix on the following content paragraph.
    claim_num = 0
    claim_headings_to_remove = []
    body = doc.element.body
    all_paras = list(body.iterchildren(qn('w:p')))

    for i, p_elem in enumerate(all_paras):
        # Check if this paragraph has a Heading2 pStyle
        pPr = p_elem.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None and pStyle.get(qn('w:val'), '').startswith('Heading'):
                # Get text content
                text_parts = []
                for r in p_elem.iterchildren(qn('w:r')):
                    for t in r.iterchildren(qn('w:t')):
                        if t.text:
                            text_parts.append(t.text)
                full_text = ''.join(text_parts).strip()

                if re.match(r'^权利要求\d+$', full_text):
                    claim_num += 1
                    claim_headings_to_remove.append(p_elem)
                    # Find the next paragraph and prefix with "N. "
                    if i + 1 < len(all_paras):
                        next_p = all_paras[i + 1]
                        # Find the first run's first text element
                        first_run = next_p.find(qn('w:r'))
                        if first_run is not None:
                            first_t = first_run.find(qn('w:t'))
                            if first_t is not None and first_t.text:
                                first_t.text = f"{claim_num}. {first_t.text}"
                            else:
                                # Insert a new text element
                                new_t = OxmlElement('w:t')
                                new_t.text = f"{claim_num}. "
                                new_t.set(qn('xml:space'), 'preserve')
                                first_run.append(new_t)
                        else:
                            # No run in next paragraph, create one
                            new_r = OxmlElement('w:r')
                            new_t = OxmlElement('w:t')
                            new_t.text = f"{claim_num}. "
                            new_t.set(qn('xml:space'), 'preserve')
                            new_r.append(new_t)
                            next_p.append(new_r)

    for p_elem in claim_headings_to_remove:
        p_elem.getparent().remove(p_elem)

    # Inline the auto-generated claim number into the paragraph text and drop
    # the numPr reference, so the paragraph follows standard body indent rules
    # (firstLineChars=200) instead of being overridden by numbering.xml's
    # abstractNum left/hanging indent.
    claim_num_inline = 0
    for p_elem in body.iterchildren(qn('w:p')):
        pPr = p_elem.find(qn('w:pPr'))
        if pPr is None:
            continue
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            continue
        first_run = p_elem.find(qn('w:r'))
        if first_run is None:
            continue
        first_t = first_run.find(qn('w:t'))
        claim_num_inline += 1
        prefix = f'{claim_num_inline}. '
        if first_t is not None and first_t.text and not re.match(r'^\d+[\.、]\s', first_t.text):
            first_t.text = prefix + first_t.text
            first_t.set(qn('xml:space'), 'preserve')
        elif first_t is None:
            new_t = OxmlElement('w:t')
            new_t.text = prefix
            new_t.set(qn('xml:space'), 'preserve')
            first_run.append(new_t)
        pPr.remove(numPr)

    if claim_num_inline > 0:
        print(f"[post-process] Inlined {claim_num_inline} claim numbers and dropped numPr", file=sys.stderr)

    if claim_num > 0:
        print(f"[post-process] Converted {claim_num} claim headings to numbered paragraphs", file=sys.stderr)

    # --- Step 4: Remove pandoc styles and apply paragraph formatting ---
    # Sub-section headings under 说明书 become bold paragraphs without first-line indent.
    spec_subsections = SPEC_SUBSECTION_TITLES
    pandoc_styles = {'Heading1', 'Heading2', 'Heading3', 'Heading4', 'Heading5',
                     'Compact', 'BodyText', 'FirstParagraph', 'SourceCode'}

    for para in doc.paragraphs:
        p_elem = para._p
        pPr = p_elem.find(qn('w:pPr'))

        # Check if paragraph contains math (m:oMath)
        has_math = p_elem.find('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath') is not None

        # Check if paragraph is empty
        para_text = para.text.strip() if para.text else ''
        if not para_text and not has_math:
            # Also check runs for any text
            run_texts = []
            for r in p_elem.iterchildren(qn('w:r')):
                for t in r.iterchildren(qn('w:t')):
                    if t.text:
                        run_texts.append(t.text)
            if not ''.join(run_texts).strip():
                continue

        # Check if paragraph is a figure caption (centered alignment)
        is_caption = False
        if pPr is not None:
            jc_elem = pPr.find(qn('w:jc'))
            if jc_elem is not None and jc_elem.get(qn('w:val')) == 'center':
                is_caption = True

        if is_caption:
            # Captions only get line spacing — no first-line indent.
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p_elem.insert(0, pPr)
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            spacing.set(qn('w:line'), '360')
            spacing.set(qn('w:lineRule'), 'auto')
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
            continue
        # Math-bearing paragraphs are still body text and must follow the same
        # firstLineChars=200 / jc=both / line=360 rules. Do not skip them.

        is_claim_paragraph = re.match(r'^[0-9]+[\.、]\s+', para_text) is not None
        is_internal_subheading = para_text in spec_subsections or INTERNAL_SUBHEADING_RE.match(para_text)

        # Remove any pandoc-generated pStyle and convert headings to bold
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                style_val = pStyle.get(qn('w:val'), '')
                if style_val in pandoc_styles or style_val.startswith('Heading'):
                    is_bold_heading = style_val in ('Heading2', 'Heading 2', 'Heading3', 'Heading 3')
                    pPr.remove(pStyle)
                    if is_bold_heading:
                        for r in p_elem.iterchildren(qn('w:r')):
                            rPr = r.find(qn('w:rPr'))
                            if rPr is None:
                                rPr = OxmlElement('w:rPr')
                                r.insert(0, rPr)
                            b = rPr.find(qn('w:b'))
                            if b is None:
                                b = OxmlElement('w:b')
                                rPr.append(b)

        # Apply paragraph formatting: line spacing, indent, justification
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p_elem.insert(0, pPr)

        # Line spacing: 360 twips (1.5x at 240 twips per line)
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:line'), '360')
        spacing.set(qn('w:lineRule'), 'auto')
        if is_internal_subheading:
            # Spec subheadings (发明名称/技术领域/.../具体实施方式/实施例X/可选实施方式)
            # need vertical breathing room from preceding text and following body
            # paragraphs so they read as headings rather than inline lines.
            spacing.set(qn('w:before'), '240')
            spacing.set(qn('w:after'), '120')
        else:
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')

        # First line indent: 200 (2-char indent)
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        if is_internal_subheading:
            ind.set(qn('w:firstLineChars'), '0')
            ind.set(qn('w:firstLine'), '0')
            if is_internal_subheading:
                for r in p_elem.iterchildren(qn('w:r')):
                    rPr = r.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        r.insert(0, rPr)
                    b = rPr.find(qn('w:b'))
                    if b is None:
                        b = OxmlElement('w:b')
                        rPr.append(b)
        else:
            ind.set(qn('w:firstLineChars'), CNIPA_FIRST_LINE_CHARS)
            ind.set(qn('w:firstLine'), CNIPA_FIRST_LINE_TWIPS)

        # Justification: both (justified)
        jc_elem = pPr.find(qn('w:jc'))
        if jc_elem is None:
            jc_elem = OxmlElement('w:jc')
            pPr.append(jc_elem)
        jc_elem.set(qn('w:val'), 'both')

        # Apply run font and size
        for r in p_elem.iterchildren(qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)

    # --- Step 5: Embed figures ---
    if figure_manifest and Path(figure_manifest).exists():
        # Auto-render mermaid-mmdc entries before checking imagePath, so the
        # caller does not need a separate manual mmdc step.
        _auto_render_mermaid(Path(figure_manifest))
        manifest = json.loads(Path(figure_manifest).read_text(encoding='utf-8'))
        manifest_dir = Path(figure_manifest).parent
        entries = []
        missing = []
        for index, entry in enumerate(manifest.get('entries', [])):
            status = entry.get('generationStatus', '')
            if status == 'skipped-with-authorization' and entry.get('userAuthorization'):
                continue
            image_path = entry.get('imagePath')
            if not image_path:
                missing.append(f'entries[{index}]: missing imagePath')
                continue
            full_path = Path(image_path)
            if not full_path.is_absolute():
                full_path = (manifest_dir / full_path).resolve()
                # Tolerate patent-root-relative imagePath (e.g. "drafts/figures/<slug>/x.png")
                # that duplicates the manifest dir prefix: fall back to the basename
                # inside the manifest directory. imagePath should be manifest-dir-relative
                # (a bare filename), but producers occasionally write the full relative path.
                if not full_path.exists():
                    alt = (manifest_dir / Path(image_path).name).resolve()
                    if alt.exists():
                        full_path = alt
            if not full_path.exists():
                missing.append(f'entries[{index}]: imagePath does not exist: {full_path}')
                continue
            normalized = dict(entry)
            normalized['imagePath'] = str(full_path)
            entries.append(normalized)
        if missing:
            print("ERROR: Figure manifest contains missing image assets:", file=sys.stderr)
            for item in missing:
                print(f"  - {item}", file=sys.stderr)
            print("Generate figure assets before export or mark entries as skipped-with-authorization.", file=sys.stderr)
            sys.exit(1)
        if entries:
            # Insert section break before figures so they get "说明书附图" header
            break_para = doc.add_paragraph()
            bp_elem = break_para._p
            bpPr = bp_elem.find(qn('w:pPr'))
            if bpPr is None:
                bpPr = OxmlElement('w:pPr')
                bp_elem.insert(0, bpPr)
            bpPr.append(make_sectPr_with_header(doc, SECTION_HEADERS.get('说明书', '说明书')))

            # Update body sectPr to use 说明书附图 header for the figures section
            body_sectPr = body.find(qn('w:sectPr'))
            if body_sectPr is not None:
                for existing_hdr in body_sectPr.findall(qn('w:headerReference')):
                    body_sectPr.remove(existing_hdr)
                apply_cnipa_section_properties(doc, body_sectPr, SECTION_HEADERS['说明书附图'])

        for entry in entries:
            full_path = Path(entry['imagePath'])
            if full_path.suffix.lower() == '.svg':
                png_path = full_path.with_suffix('.png')
                try:
                    import cairosvg
                    cairosvg.svg2png(url=str(full_path), write_to=str(png_path), output_width=1200)
                    full_path = png_path
                except ImportError:
                    print(f"WARNING: cairosvg not available, skipping SVG figure: {entry['imagePath']}", file=sys.stderr)
                    continue
            doc.add_paragraph()
            doc.add_picture(str(full_path), width=Inches(5.5))
            figure_label = normalize_figure_number_label(entry)
            if figure_label:
                lbl_para = doc.add_paragraph(figure_label)
                lbl_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lbl_pPr = lbl_para._p.find(qn('w:pPr'))
                if lbl_pPr is None:
                    lbl_pPr = OxmlElement('w:pPr')
                    lbl_para._p.insert(0, lbl_pPr)
                lbl_spacing = lbl_pPr.find(qn('w:spacing'))
                if lbl_spacing is None:
                    lbl_spacing = OxmlElement('w:spacing')
                    lbl_pPr.append(lbl_spacing)
                lbl_spacing.set(qn('w:line'), '360')
                lbl_spacing.set(qn('w:lineRule'), 'auto')
                lbl_spacing.set(qn('w:before'), '0')
                lbl_spacing.set(qn('w:after'), '0')

    # --- Step 5b: Remove "附图交接清单" table from document ---
    body = doc.element.body
    removing = False
    to_remove = []
    for p_elem in list(body.iterchildren()):
        tag = p_elem.tag.split('}')[-1] if '}' in p_elem.tag else p_elem.tag
        if tag == 'p':
            texts = [t.text for t in p_elem.iter(qn('w:t')) if t.text]
            text = ''.join(texts).strip()
            if text.startswith('附图交接清单'):
                removing = True
                to_remove.append(p_elem)
            elif removing:
                if text and (text.startswith('##') or text.startswith('#')):
                    removing = False
                else:
                    to_remove.append(p_elem)
        elif tag == 'tbl' and removing:
            to_remove.append(p_elem)
    for elem in to_remove:
        body.remove(elem)

    # --- Step 6: Set page layout on body sectPr ---
    # A4 page size and margins matching cn-patent-draft-v4.docx template
    body = doc.element.body
    body_sectPr = body.find(qn('w:sectPr'))
    if body_sectPr is None:
        body_sectPr = OxmlElement('w:sectPr')
        body.append(body_sectPr)

    apply_cnipa_section_properties(doc, body_sectPr)

    # Official specification page starts with the invention title centered;
    # keep the title text but remove the internal "发明名称" label.
    remove_paras = []
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip() == '发明名称':
            remove_paras.append(para._p)
            for next_para in doc.paragraphs[idx + 1:]:
                if next_para.text.strip():
                    next_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    next_pPr = next_para._p.find(qn('w:pPr'))
                    if next_pPr is None:
                        next_pPr = OxmlElement('w:pPr')
                        next_para._p.insert(0, next_pPr)
                    ind = next_pPr.find(qn('w:ind'))
                    if ind is None:
                        ind = OxmlElement('w:ind')
                        next_pPr.append(ind)
                    ind.set(qn('w:firstLineChars'), '0')
                    ind.set(qn('w:firstLine'), '0')
                    for run in next_para.runs:
                        run.bold = True
                    break
    for p_elem in remove_paras:
        p_elem.getparent().remove(p_elem)

    _patch_omath_cambria_math(doc)

    doc.save(str(docx_path))
    _prune_orphan_header_footer_parts(docx_path)


def _patch_omath_cambria_math(doc):
    """Ensure every <m:r> inside <m:oMath> carries Cambria Math font.

    Pandoc-produced inline math (from $...$ or \(...\)) emits <m:oMath>
    runs without <w:rPr><w:rFonts ascii="Cambria Math".../></w:rPr>. Word
    for Mac requires this font marker to invoke its OMML renderer; without
    it the formula appears blank or as raw text. The python-docx path for
    display math (_latex_to_omml) already injects this; we replicate the
    same injection here for everything pandoc produced so cross-platform
    rendering matches reference.docx.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    body = doc.element.body
    for m_r in body.iter(f'{{{M_NS}}}r'):
        wrPr = m_r.find(qn('w:rPr'))
        if wrPr is None:
            wrPr = OxmlElement('w:rPr')
            m_r.insert(0, wrPr)
        rFonts = wrPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            wrPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), 'Cambria Math')
        rFonts.set(qn('w:hAnsi'), 'Cambria Math')


def _prune_orphan_header_footer_parts(docx_path):
    """Strip header/footer parts that pandoc copied from reference.docx but
    are not referenced by any section. Also drops their _rels and the
    [Content_Types].xml Override entries.
    """
    import zipfile, re as _re, shutil as _shutil
    docx_path = Path(docx_path)
    with zipfile.ZipFile(docx_path, 'r') as zf:
        names = zf.namelist()
        if 'word/document.xml' not in names or 'word/_rels/document.xml.rels' not in names:
            return
        document_xml = zf.read('word/document.xml').decode('utf-8')
        document_rels = zf.read('word/_rels/document.xml.rels').decode('utf-8')

    referenced_rids = set(_re.findall(
        r'(?:headerReference|footerReference)[^>]*r:id="([^"]+)"', document_xml))

    rid_to_target = {}
    for m in _re.finditer(r'<Relationship([^>]*?)/?>', document_rels):
        attrs = dict(_re.findall(r'(\w+)="([^"]+)"', m.group(1)))
        if attrs.get('Type', '').endswith(('/header', '/footer')):
            rid_to_target[attrs.get('Id')] = attrs.get('Target')

    orphan_rids, orphan_files = set(), set()
    for rid, target in rid_to_target.items():
        target_path = target if target.startswith('word/') else f'word/{target}'
        if rid not in referenced_rids:
            orphan_rids.add(rid)
            orphan_files.add(target_path)

    if not orphan_rids:
        return

    orphan_zip_paths = set(orphan_files)
    for fpath in list(orphan_files):
        rels_name = fpath.replace('word/', 'word/_rels/') + '.rels'
        if rels_name in names:
            orphan_zip_paths.add(rels_name)

    new_rels = document_rels
    for rid in orphan_rids:
        new_rels = _re.sub(
            rf'<Relationship[^>]*?Id="{_re.escape(rid)}"[^>]*?/?>', '', new_rels)

    tmp = docx_path.with_suffix('.tmp.docx')
    with zipfile.ZipFile(docx_path, 'r') as zin, \
         zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename in orphan_zip_paths:
                continue
            data = zin.read(item.filename)
            if item.filename == 'word/_rels/document.xml.rels':
                data = new_rels.encode('utf-8')
            elif item.filename == '[Content_Types].xml':
                ct = data.decode('utf-8')
                for fpath in orphan_files:
                    part_name = '/' + fpath
                    ct = _re.sub(
                        rf'<Override[^/]*PartName="{_re.escape(part_name)}"[^/]*/>',
                        '', ct)
                data = ct.encode('utf-8')
            zout.writestr(item, data)
    _shutil.move(tmp, docx_path)


def _auto_render_mermaid(manifest_path: Path) -> None:
    """Best-effort: render any mermaid-mmdc entries via render-mermaid-figures.py.

    Idempotent (the helper itself skips up-to-date PNGs). Soft-fails: if the
    helper or mmdc is unavailable, the downstream image-existence check still
    surfaces missing assets with a clear error.
    """
    helper = Path(__file__).resolve().parent / 'render-mermaid-figures.py'
    if not helper.is_file():
        return
    try:
        result = subprocess.run(
            [sys.executable, str(helper), '--manifest', str(manifest_path)],
            capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=300)
    except (subprocess.TimeoutExpired, FileNotFoundError):
        return
    for line in (result.stdout or '').splitlines():
        if line.strip():
            print(line, file=sys.stderr)
    if result.returncode != 0:
        for line in (result.stderr or '').splitlines():
            if line.strip():
                print(line, file=sys.stderr)


def _md_to_docx_python(input_path, output_path):
    """Parse Markdown and build a structured DOCX with Heading styles and math."""
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    text = Path(input_path).read_text(encoding='utf-8')
    text = truncate_non_patent_content(text)

    doc = Document()

    # Set docDefaults to match reference.docx
    styles_elem = doc.styles.element
    docDefaults = styles_elem.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = OxmlElement('w:docDefaults')
        styles_elem.insert(0, docDefaults)
    rPrDefault = docDefaults.find(qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = OxmlElement('w:rPrDefault')
        docDefaults.append(rPrDefault)
    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        rPrDefault.append(rPr)
    # Clear and set fonts
    for child in list(rPr):
        rPr.remove(child)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), '宋体')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '28')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '28')
    rPr.append(szCs)
    # Remove pPrDefault (python-docx template adds spacing after=200, line=276)
    pPrDefault = docDefaults.find(qn('w:pPrDefault'))
    if pPrDefault is not None:
        docDefaults.remove(pPrDefault)

    heading_re = re.compile(r'^(#{1,4})\s+(.*)')
    latex_fence_start = re.compile(r'^```latex\s*$')
    fence_end = re.compile(r'^```\s*$')
    inline_math_re = re.compile(r'\\\((.*?)\\\)')
    mermaid_fence_start = re.compile(r'^```mermaid\s*$')

    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip mermaid blocks
        if mermaid_fence_start.match(stripped):
            i += 1
            while i < len(lines) and not fence_end.match(lines[i].strip()):
                i += 1
            i += 1
            continue

        # LaTeX display math block
        if latex_fence_start.match(stripped):
            i += 1
            latex_buf = []
            while i < len(lines) and not fence_end.match(lines[i].strip()):
                latex_buf.append(lines[i])
                i += 1
            i += 1
            # Insert as oMathPara
            para = doc.add_paragraph()
            p_elem = para._p
            latex_str = '\n'.join(latex_buf).strip()
            _insert_omath(p_elem, latex_str)
            continue

        # Heading
        m = heading_re.match(stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            doc.add_heading(heading_text, level=level)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Normal paragraph (may contain inline math)
        para = doc.add_paragraph()
        _add_runs_with_inline_math(para, stripped, inline_math_re)
        i += 1

    doc.save(str(output_path))


def _latex_to_omml(latex_str):
    """Convert LaTeX string to OMML element via MathML + Word XSLT, with Cambria Math font."""
    from lxml import etree
    import latex2mathml.converter

    xslt_path = None
    if platform.system() == "Windows":
        candidates = [
            Path(r'C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL'),
            Path(r'C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL'),
        ]
        for c in candidates:
            if c.exists():
                xslt_path = c
                break
        if not xslt_path:
            for p in Path(r'C:\Program Files').rglob('MML2OMML.XSL'):
                xslt_path = p
                break
    else:
        bundled = Path(__file__).parent / 'MML2OMML.XSL'
        if bundled.exists():
            xslt_path = bundled
        else:
            for search_dir in [Path('/usr/share/docx2tex'), Path('/usr/local/share'),
                               Path.home() / '.local' / 'share']:
                for p in search_dir.rglob('MML2OMML.XSL'):
                    xslt_path = p
                    break
                if xslt_path:
                    break

    if not xslt_path or not xslt_path.exists():
        raise FileNotFoundError("MML2OMML.XSL not found. On Linux, place it in the scripts/ directory.")

    mathml = latex2mathml.converter.convert(latex_str)
    mathml_tree = etree.fromstring(mathml.encode('utf-8'))
    xslt = etree.parse(str(xslt_path))
    transform = etree.XSLT(xslt)
    omml = transform(mathml_tree)
    root = omml.getroot()

    # Add Cambria Math font to all m:r elements (matching reference.docx)
    M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for mr in root.iter(f'{{{M_NS}}}r'):
        # Find or create w:rPr inside m:r
        wrPr = mr.find(f'{{{W_NS}}}rPr')
        if wrPr is None:
            wrPr = etree.SubElement(mr, f'{{{W_NS}}}rPr')
            mr.insert(0, wrPr)
        rFonts = wrPr.find(f'{{{W_NS}}}rFonts')
        if rFonts is None:
            rFonts = etree.SubElement(wrPr, f'{{{W_NS}}}rFonts')
        rFonts.set(f'{{{W_NS}}}ascii', 'Cambria Math')
        rFonts.set(f'{{{W_NS}}}hAnsi', 'Cambria Math')

    return root


def _insert_omath(p_elem, latex_str):
    """Insert a display math (oMathPara) element converted from LaTeX."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Add pPr to match reference.docx formatting for formula paragraphs
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '360')
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:firstLineChars'), '200')
    ind.set(qn('w:firstLine'), '560')
    pPr.append(ind)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    pPr.append(jc)

    try:
        omml_elem = _latex_to_omml(latex_str)
        oMathPara = OxmlElement('m:oMathPara')
        oMathPara.append(omml_elem)
        p_elem.append(oMathPara)
    except Exception:
        oMathPara = OxmlElement('m:oMathPara')
        oMath = OxmlElement('m:oMath')
        mr = OxmlElement('m:r')
        mt = OxmlElement('m:t')
        mt.text = latex_str
        mr.append(mt)
        oMath.append(mr)
        oMathPara.append(oMath)
        p_elem.append(oMathPara)


def _add_runs_with_inline_math(para, text, inline_math_re):
    """Add runs to paragraph, converting \\(...\\) to oMath and **...** to bold."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    bold_re = re.compile(r'\*\*(.*?)\*\*')
    math_parts = inline_math_re.split(text)
    for idx, part in enumerate(math_parts):
        if idx % 2 == 0:
            if part:
                bold_parts = bold_re.split(part)
                for bidx, bpart in enumerate(bold_parts):
                    if not bpart:
                        continue
                    if bidx % 2 == 0:
                        para.add_run(bpart)
                    else:
                        para.add_run(bpart).bold = True
        else:
            p_elem = para._p
            try:
                omml_elem = _latex_to_omml(part)
                p_elem.append(omml_elem)
            except Exception:
                oMath = OxmlElement('m:oMath')
                mr = OxmlElement('m:r')
                mt = OxmlElement('m:t')
                mt.text = part
                mr.append(mt)
                oMath.append(mr)
                p_elem.append(oMath)


def export_with_word_com(input_path, output_path):
    """Windows-only: build structured DOCX from Markdown, then normalize via Word COM."""
    if platform.system() != "Windows":
        return False
    try:
        import win32com.client
    except ImportError:
        return False

    try:
        # Step 1: Build structured DOCX with python-docx
        _md_to_docx_python(input_path, output_path)

        # Step 2: Open in Word COM and re-save to normalize
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(Path(output_path).resolve()))
        doc.SaveAs2(str(Path(output_path).resolve()), FileFormat=16)
        doc.Close()
        word.Quit()
        return True
    except Exception as e:
        print(f"WARNING: Word COM normalization failed ({e}), using python-docx output directly", file=sys.stderr)
        # If Word COM fails but python-docx succeeded, still return True
        if Path(output_path).exists():
            return True
        return False


def verify_section_headers(docx_path, has_figures=False):
    """Verify all required section headers exist in the exported DOCX."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        return

    doc = Document(str(docx_path))
    found_headers = set()
    for rel in doc.part.rels.values():
        if 'header' in rel.reltype:
            hdr_elem = rel.target_part.element
            for p in hdr_elem.iterchildren(qn('w:p')):
                texts = [t.text for t in p.iter(qn('w:t')) if t.text]
                found_headers.add(''.join(texts))

    required = {'说明书摘要', '权利要求书', '说明书'}
    if has_figures:
        required.add('说明书附图')

    missing = required - found_headers
    if missing:
        short_names = [h.replace(' ', '') for h in missing]
        print(f"WARNING: Missing section headers: {', '.join(short_names)}", file=sys.stderr)
    else:
        expected = '摘要/权利要求书/说明书'
        if has_figures:
            expected += '/说明书附图'
        print(f"[verify] Section headers OK: {expected}", file=sys.stderr)


SPLIT_FILE_NAMES = {
    '说明书摘要': '说明书摘要.docx',
    '权利要求书': '权利要求书.docx',
    '说明书': '说明书.docx',
    '说明书附图': '说明书附图.docx',
}
SPLIT_HEADER_GROUPS = {
    '摘要附图': '说明书摘要',
}


def _header_text_for_sectpr(doc, sectPr):
    """Resolve the header text referenced by a section properties element."""
    from docx.oxml.ns import qn

    header_ref = sectPr.find(qn('w:headerReference'))
    if header_ref is None:
        return ''
    rid = header_ref.get(qn('r:id'))
    if not rid or rid not in doc.part.rels:
        return ''
    header_part = doc.part.rels[rid].target_part
    texts = [t.text for t in header_part.element.iter(qn('w:t')) if t.text]
    return ''.join(texts).strip()


def _strip_section_properties(element):
    """Remove embedded section properties before placing content into a split file."""
    from docx.oxml.ns import qn

    pPr = element.find(qn('w:pPr'))
    if pPr is None:
        return
    for sectPr in list(pPr.findall(qn('w:sectPr'))):
        pPr.remove(sectPr)


def _clear_document_body(doc):
    body = doc.element.body
    for child in list(body):
        body.remove(child)
    return body


def split_docx_by_sections(docx_path, split_output_dir):
    """Create per-module DOCX files from the already formatted full DOCX."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    source = Document(str(docx_path))
    split_dir = Path(split_output_dir)
    split_dir.mkdir(parents=True, exist_ok=True)

    body = source.element.body
    current_elements = []
    sections = []

    for elem in body:
        if elem.tag == qn('w:p'):
            current_elements.append(elem)
            pPr = elem.find(qn('w:pPr'))
            sectPr = pPr.find(qn('w:sectPr')) if pPr is not None else None
            if sectPr is not None:
                header_text = _header_text_for_sectpr(source, sectPr)
                sections.append((header_text, list(current_elements)))
                current_elements = []
        elif elem.tag == qn('w:sectPr'):
            header_text = _header_text_for_sectpr(source, elem)
            sections.append((header_text, list(current_elements)))
            current_elements = []
        else:
            current_elements.append(elem)

    written = []
    grouped_sections = []
    group_index = {}
    for header_text, elements in sections:
        group_header = SPLIT_HEADER_GROUPS.get(header_text, header_text)
        if group_header in group_index:
            grouped_sections[group_index[group_header]][1].extend(elements)
        else:
            group_index[group_header] = len(grouped_sections)
            grouped_sections.append((group_header, list(elements)))

    for header_text, elements in grouped_sections:
        file_name = SPLIT_FILE_NAMES.get(header_text)
        if not file_name:
            continue

        has_content = False
        prepared_elements = []
        for elem in elements:
            copied = copy.deepcopy(elem)
            _strip_section_properties(copied)
            texts = [t.text for t in copied.iter(qn('w:t')) if t.text]
            if ''.join(texts).strip() or copied.tag != qn('w:p'):
                has_content = True
            prepared_elements.append(copied)

        if not has_content:
            continue

        out_path = split_dir / file_name
        # Copy the whole OPC package first so drawings keep their image relationships.
        shutil.copyfile(docx_path, out_path)
        target = Document(str(out_path))
        target_body = _clear_document_body(target)
        for copied in prepared_elements:
            target_body.append(copied)

        sectPr = OxmlElement('w:sectPr')
        apply_cnipa_section_properties(target, sectPr, header_text)
        target_body.append(sectPr)
        target.save(str(out_path))
        written.append(out_path)

    return written


def main():
    parser = argparse.ArgumentParser(description='Export patent Markdown to DOCX')
    parser.add_argument('--input', required=True, help='Input Markdown path')
    parser.add_argument('--output', default='', help='Output DOCX path')
    parser.add_argument('--reference-doc', default='', help='Pandoc reference.docx template')
    parser.add_argument('--figure-manifest', default='', help='Figure manifest JSON path')
    parser.add_argument('--render-equations', action='store_true', help='Enable equation rendering')
    parser.add_argument('--state-path', default='', help='Patent iteration state JSON path')
    parser.add_argument('--skip-gate-check', action='store_true', help='Skip Gate C status check (requires explicit user authorization)')
    parser.add_argument('--split-output-dir', default='', help='Optional directory for per-module CNIPA split DOCX files')
    args = parser.parse_args()

    # Gate C hard check: refuse export if gate_c.status != approved
    if not args.skip_gate_check:
        state_path = Path(args.state_path) if args.state_path else Path(args.input).parent.parent.parent / 'state' / 'patent-iteration-state.json'
        if state_path.exists():
            try:
                state = json.loads(state_path.read_text(encoding='utf-8'))
                gate_c_status = state.get('gate_c', {}).get('status', '')
                if gate_c_status not in ('approved', 'confirmed', 'passed'):
                    print(f"ERROR: Gate C not approved (status='{gate_c_status}'). "
                          f"Cannot export DOCX before Gate C confirmation. "
                          f"State file: {state_path}", file=sys.stderr)
                    sys.exit(1)
            except (json.JSONDecodeError, KeyError) as e:
                print(f"WARNING: Cannot read Gate C status from {state_path}: {e}", file=sys.stderr)

    preflight_check()

    input_path = Path(args.input).resolve()
    if not input_path.exists():
        print(f"ERROR: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    validate_markdown_source(input_path)

    output_path = Path(args.output) if args.output else input_path.with_suffix('.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)

    figure_manifest = args.figure_manifest

    used_word_com = False
    if platform.system() == "Windows" and not shutil.which("pandoc"):
        used_word_com = export_with_word_com(input_path, output_path)

    if not used_word_com:
        check_pandoc()
        reference_doc = args.reference_doc if args.reference_doc else None
        if not reference_doc:
            script_dir = Path(__file__).parent
            default_ref = script_dir.parent / 'assets' / 'reference.docx'
            if default_ref.exists():
                reference_doc = str(default_ref)
        export_with_pandoc(input_path, output_path, reference_doc)

    post_process_docx(output_path, figure_manifest if figure_manifest else None)
    verify_section_headers(output_path, has_figures=bool(figure_manifest))
    if args.split_output_dir:
        split_paths = split_docx_by_sections(output_path, args.split_output_dir)
        if split_paths:
            for sp in split_paths:
                _prune_orphan_header_footer_parts(sp)
            names = ', '.join(path.name for path in split_paths)
            print(f"[split-export] Generated split DOCX files: {names}", file=sys.stderr)
        else:
            print("[split-export] WARNING: no split DOCX files generated", file=sys.stderr)
    backend = "Word COM + post-process" if used_word_com else "pandoc + post-process"
    print(f"[export-patent-draft-docx] Exported via {backend}: {output_path}")


if __name__ == "__main__":
    for _stream in (sys.stdout, sys.stderr):
        if hasattr(_stream, "reconfigure"):
            _stream.reconfigure(encoding="utf-8")
    main()
