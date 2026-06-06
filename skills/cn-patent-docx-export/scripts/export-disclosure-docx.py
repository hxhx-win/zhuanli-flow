#!/usr/bin/env python3
"""将技术交底书 Markdown 导出为 DOCX（pandoc 公式 OMML + CNIPA 正文排版）。"""
import argparse
import importlib.util
import re
import subprocess
import sys
import tempfile
from pathlib import Path

_SCRIPT_DIR = Path(__file__).resolve().parent
_PATENT_EXPORT = _SCRIPT_DIR / 'export-patent-draft-docx.py'

_spec = importlib.util.spec_from_file_location('patent_export', _PATENT_EXPORT)
_pe = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(_pe)

CNIPA_PAGE_SIZE = _pe.CNIPA_PAGE_SIZE
CNIPA_PAGE_MARGINS = _pe.CNIPA_PAGE_MARGINS
CNIPA_BODY_FONT_EAST_ASIA = _pe.CNIPA_BODY_FONT_EAST_ASIA
CNIPA_BODY_FONT_ASCII = _pe.CNIPA_BODY_FONT_ASCII
CNIPA_BODY_FONT_SIZE = _pe.CNIPA_BODY_FONT_SIZE
CNIPA_LINE_SPACING = _pe.CNIPA_LINE_SPACING
CNIPA_FIRST_LINE_CHARS = _pe.CNIPA_FIRST_LINE_CHARS
CNIPA_FIRST_LINE_TWIPS = _pe.CNIPA_FIRST_LINE_TWIPS
preprocess_latex_fences = _pe.preprocess_latex_fences
apply_cnipa_section_properties = _pe.apply_cnipa_section_properties
make_footer_part = _pe.make_footer_part


def export_disclosure_with_pandoc(input_path: Path, output_path: Path, reference_doc=None):
    """Pandoc 导出，保留全文并启用 LaTeX 公式。"""
    text = input_path.read_text(encoding='utf-8')
    processed = preprocess_latex_fences(text)
    processed = re.sub(r'\\\((.+?)\\\)', r'$\1$', processed)

    tmp = tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8')
    tmp.write(processed)
    tmp.close()

    cmd = [
        'pandoc', tmp.name, '-o', str(output_path),
        '--from', 'markdown+tex_math_dollars+raw_tex',
        '--wrap', 'none',
    ]
    if reference_doc and Path(reference_doc).exists():
        cmd.extend(['--reference-doc', str(reference_doc)])
    result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8', errors='replace')
    Path(tmp.name).unlink(missing_ok=True)
    if result.returncode != 0:
        print(f'ERROR: pandoc failed: {result.stderr}', file=sys.stderr)
        sys.exit(1)


def _get_para_style(p_elem, qn):
    p_pr = p_elem.find(qn('w:pPr'))
    if p_pr is None:
        return ''
    p_style = p_pr.find(qn('w:pStyle'))
    if p_style is None:
        return ''
    return p_style.get(qn('w:val'), '')


def _set_run_fonts(r_pr, qn, oxml_element, east_asia=None, ascii_font=None, size=None, mono=False):
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = oxml_element('w:rFonts')
        r_pr.append(r_fonts)
    if mono:
        r_fonts.set(qn('w:ascii'), 'Courier New')
        r_fonts.set(qn('w:hAnsi'), 'Courier New')
        r_fonts.set(qn('w:eastAsia'), '宋体')
    else:
        r_fonts.set(qn('w:ascii'), ascii_font or CNIPA_BODY_FONT_ASCII)
        r_fonts.set(qn('w:hAnsi'), ascii_font or CNIPA_BODY_FONT_ASCII)
        r_fonts.set(qn('w:eastAsia'), east_asia or CNIPA_BODY_FONT_EAST_ASIA)
    if size:
        sz = r_pr.find(qn('w:sz'))
        if sz is None:
            sz = oxml_element('w:sz')
            r_pr.append(sz)
        sz.set(qn('w:val'), size)


def post_process_disclosure_docx(docx_path: Path):
    """交底书版式后处理：页边距、页码、标题层级、正文/代码/公式段落格式。"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    oxml = OxmlElement
    doc = Document(str(docx_path))
    if not doc.paragraphs:
        print('WARNING: empty DOCX', file=sys.stderr)
        return

    pandoc_styles = {
        'Heading1', 'Heading2', 'Heading3', 'Heading4', 'Heading5',
        'Compact', 'BodyText', 'FirstParagraph', 'SourceCode', 'BlockText',
    }

    for para in doc.paragraphs:
        p_elem = para._p
        p_pr = p_elem.find(qn('w:pPr'))
        style_val = _get_para_style(p_elem, qn)
        has_math = p_elem.find(
            './/{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath'
        ) is not None
        para_text = (para.text or '').strip()

        if not para_text and not has_math:
            run_texts = []
            for r in p_elem.iterchildren(qn('w:r')):
                for t in r.iterchildren(qn('w:t')):
                    if t.text:
                        run_texts.append(t.text)
            if not ''.join(run_texts).strip() and not has_math:
                continue

        is_source_code = style_val == 'SourceCode'
        is_heading = style_val.startswith('Heading') if style_val else False
        is_block_meta = style_val in ('BlockText', 'Compact')
        is_caption = False
        if p_pr is not None:
            jc_elem = p_pr.find(qn('w:jc'))
            if jc_elem is not None and jc_elem.get(qn('w:val')) == 'center':
                is_caption = True

        if p_pr is None:
            p_pr = OxmlElement('w:pPr')
            p_elem.insert(0, p_pr)

        # 去掉 pandoc 样式名
        p_style = p_pr.find(qn('w:pStyle'))
        if p_style is not None and p_style.get(qn('w:val'), '') in pandoc_styles:
            p_pr.remove(p_style)

        spacing = p_pr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            p_pr.append(spacing)
        spacing.set(qn('w:line'), CNIPA_LINE_SPACING)
        spacing.set(qn('w:lineRule'), 'auto')

        ind = p_pr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            p_pr.append(ind)

        jc_elem = p_pr.find(qn('w:jc'))
        if jc_elem is None:
            jc_elem = OxmlElement('w:jc')
            p_pr.append(jc_elem)

        if style_val == 'Heading1':
            spacing.set(qn('w:before'), '360')
            spacing.set(qn('w:after'), '240')
            ind.set(qn('w:firstLineChars'), '0')
            ind.set(qn('w:firstLine'), '0')
            jc_elem.set(qn('w:val'), 'center')
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p_elem.iterchildren(qn('w:r')):
                r_pr = r.find(qn('w:rPr'))
                if r_pr is None:
                    r_pr = OxmlElement('w:rPr')
                    r.insert(0, r_pr)
                b = r_pr.find(qn('w:b'))
                if b is None:
                    r_pr.append(OxmlElement('w:b'))
                _set_run_fonts(r_pr, qn, oxml, size='32')
        elif style_val in ('Heading2', 'Heading3', 'Heading4'):
            spacing.set(qn('w:before'), '240' if style_val == 'Heading2' else '120')
            spacing.set(qn('w:after'), '120')
            ind.set(qn('w:firstLineChars'), '0')
            ind.set(qn('w:firstLine'), '0')
            jc_elem.set(qn('w:val'), 'left')
            for r in p_elem.iterchildren(qn('w:r')):
                r_pr = r.find(qn('w:rPr'))
                if r_pr is None:
                    r_pr = OxmlElement('w:rPr')
                    r.insert(0, r_pr)
                b = r_pr.find(qn('w:b'))
                if b is None:
                    r_pr.append(OxmlElement('w:b'))
                _set_run_fonts(r_pr, qn, oxml, size='28' if style_val == 'Heading2' else CNIPA_BODY_FONT_SIZE)
        elif is_source_code:
            spacing.set(qn('w:before'), '60')
            spacing.set(qn('w:after'), '60')
            spacing.set(qn('w:line'), '240')
            ind.set(qn('w:firstLineChars'), '0')
            ind.set(qn('w:firstLine'), '0')
            ind.set(qn('w:left'), '360')
            jc_elem.set(qn('w:val'), 'left')
            for r in p_elem.iterchildren(qn('w:r')):
                r_pr = r.find(qn('w:rPr'))
                if r_pr is None:
                    r_pr = OxmlElement('w:rPr')
                    r.insert(0, r_pr)
                _set_run_fonts(r_pr, qn, oxml, mono=True, size='20')
        elif is_block_meta:
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '60')
            ind.set(qn('w:firstLineChars'), '0')
            ind.set(qn('w:firstLine'), '0')
            ind.set(qn('w:left'), '400')
            jc_elem.set(qn('w:val'), 'both')
            for r in p_elem.iterchildren(qn('w:r')):
                r_pr = r.find(qn('w:rPr'))
                if r_pr is None:
                    r_pr = OxmlElement('w:rPr')
                    r.insert(0, r_pr)
                _set_run_fonts(r_pr, qn, oxml, size='24')
        elif has_math and not para_text:
            spacing.set(qn('w:before'), '60')
            spacing.set(qn('w:after'), '60')
            ind.set(qn('w:firstLineChars'), '0')
            ind.set(qn('w:firstLine'), '0')
            jc_elem.set(qn('w:val'), 'center')
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif is_caption:
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
            ind.set(qn('w:firstLineChars'), '0')
            ind.set(qn('w:firstLine'), '0')
        else:
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
            ind.set(qn('w:firstLineChars'), CNIPA_FIRST_LINE_CHARS)
            ind.set(qn('w:firstLine'), CNIPA_FIRST_LINE_TWIPS)
            jc_elem.set(qn('w:val'), 'both')
            for r in p_elem.iterchildren(qn('w:r')):
                r_pr = r.find(qn('w:rPr'))
                if r_pr is None:
                    r_pr = OxmlElement('w:rPr')
                    r.insert(0, r_pr)
                _set_run_fonts(r_pr, qn, oxml, size=CNIPA_BODY_FONT_SIZE)

    # 表格：宋体 10.5pt
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    cp = para._p
                    cpr = cp.find(qn('w:pPr'))
                    if cpr is None:
                        cpr = OxmlElement('w:pPr')
                        cp.insert(0, cpr)
                    cjc = cpr.find(qn('w:jc'))
                    if cjc is None:
                        cjc = OxmlElement('w:jc')
                        cpr.append(cjc)
                    cjc.set(qn('w:val'), 'center')
                    for r in cp.iterchildren(qn('w:r')):
                        r_pr = r.find(qn('w:rPr'))
                        if r_pr is None:
                            r_pr = OxmlElement('w:rPr')
                            r.insert(0, r_pr)
                        _set_run_fonts(r_pr, qn, oxml, size='21')

    body = doc.element.body
    body_sect_pr = body.find(qn('w:sectPr'))
    if body_sect_pr is None:
        body_sect_pr = OxmlElement('w:sectPr')
        body.append(body_sect_pr)
    apply_cnipa_section_properties(doc, body_sect_pr, header_text='技术交底书')
    doc.save(str(docx_path))


def verify_disclosure_docx(docx_path: Path):
    from docx import Document

    doc = Document(str(docx_path))
    omath = 0
    for p in doc.element.body.iter():
        if p.tag.endswith('oMath'):
            omath += 1
    print(f'[verify] paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} oMath={omath}',
          file=sys.stderr)
    if omath == 0:
        print('[verify] WARN: no OMML formulas found', file=sys.stderr)


def main():
    parser = argparse.ArgumentParser(description='Export technical disclosure Markdown to DOCX')
    parser.add_argument('--input', required=True)
    parser.add_argument('--output', default='')
    parser.add_argument('--reference-doc', default='')
    args = parser.parse_args()

    _pe.preflight_check()

    input_path = Path(args.input).resolve()
    if not input_path.exists():
        print(f'ERROR: not found: {input_path}', file=sys.stderr)
        sys.exit(1)

    output_path = Path(args.output) if args.output else input_path.with_suffix('.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)

    reference_doc = args.reference_doc or None
    if not reference_doc:
        default_ref = _SCRIPT_DIR.parent / 'assets' / 'reference.docx'
        if default_ref.exists():
            reference_doc = str(default_ref)

    export_disclosure_with_pandoc(input_path, output_path, reference_doc)
    post_process_disclosure_docx(output_path)
    verify_disclosure_docx(output_path)
    print(f'[export-disclosure-docx] Exported: {output_path}')


if __name__ == '__main__':
    for stream in (sys.stdout, sys.stderr):
        if hasattr(stream, 'reconfigure'):
            stream.reconfigure(encoding='utf-8')
    main()
